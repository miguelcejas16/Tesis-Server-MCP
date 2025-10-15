from datetime import date
from typing import Optional, Literal, List
import asyncpg

'''
 * Crea un registro de reintegro temporal en la base de datos.
 * Parámetros:
 *   connection (asyncpg.Connection) — Conexión a la base de datos.
 *   afiliado_id (int) — ID del afiliado que solicita el reintegro.
 *   total_presentado (float) — Monto total presentado inicialmente.
 * Retorna:
 *   int — El ID del reintegro creado.
 * Notas:
 *   El estado inicial del reintegro se establece como 'pendiente'.
'''
async def create_temp_reintegro(
        connection: asyncpg.Connection, 
        afiliado_id: int,
        total_presentado: float,
        ) -> int: 
    try:
        query = """
            INSERT INTO public.reintegro (afiliado_id, estado, total_presentado, total_aprobado)
            VALUES ($1, 'pendiente', $2, 0.0)
            RETURNING reintegro_id
        """
        result = await connection.fetchrow(query, afiliado_id, total_presentado)

        if result:
            return result['reintegro_id']
        else:
            raise Exception("No se pudo crear el reintegro y obtener el ID.")
    except Exception as e:
        raise Exception(f"Error en utils.create_temp_reintegro: {e}")

'''
 * Agrega un ítem (práctica o medicamento) a un reintegro existente.
 * Parámetros:
 *   connection (asyncpg.Connection) — Conexión a la base de datos.
 *   reintegro_id (int) — ID del reintegro al que se agrega el ítem.
 *   tipo (Literal["M", "P"]) — Tipo de ítem ('M' para medicamento, 'P' para práctica).
 *   practica_id (Optional[int]) — ID de la práctica (si aplica).
 *   medicamento_id (Optional[int]) — ID del medicamento (si aplica).
 *   fecha_prestacion (date) — Fecha en que se realizó la prestación.
 *   monto_presentado (float) — Monto presentado para este ítem.
 * Retorna:
 *   int — El ID del ítem de reintegro creado.
'''
async def add_item_to_reintegro(
        connection: asyncpg.Connection, 
        reintegro_id: int,
        tipo: Literal["M", "P"],
        practica_id: Optional[int],
        medicamento_id: Optional[int],
        fecha_prestacion: date,
        monto_presentado: float,
        ) -> int:
    try:
        # Mapear 'M'/'P' a los valores de la base de datos y validar
        if tipo == "M":
            tipo_db = "medicamento"
            if not medicamento_id:
                raise ValueError("medicamento_id es requerido cuando el tipo es 'M'")
        elif tipo == "P":
            tipo_db = "practica"
            if not practica_id:
                raise ValueError("practica_id es requerido cuando el tipo es 'P'")
        else:
            raise ValueError("El valor de 'tipo' debe ser 'M' o 'P'")

        query = """
            INSERT INTO public.reintegro_item (
                reintegro_id, tipo, practica_id, medicamento_id, 
                fecha_prestacion, monto_presentado
            )
            VALUES ($1, $2, $3, $4, $5, $6)
            RETURNING item_id
        """
        result = await connection.fetchrow(
            query,
            reintegro_id, tipo_db, practica_id, medicamento_id,
            fecha_prestacion, monto_presentado
        )

        if result:
            return result['item_id']
        else:
            raise Exception("No se pudo agregar el ítem y obtener el ID.")
    except Exception as e:
        raise Exception(f"Error en utils.add_item_to_reintegro: {e}")

async def add_docs_reintegro(connection: asyncpg.Connection, reintegro_id: int) -> bool:
    '''
     * Marca un reintegro como esperando adjuntos y resetea la confirmación de adjuntos.
     *
     * Parámetros:
     *   connection (asyncpg.Connection) — Conexión a la base de datos.
     *   reintegro_id (int) — ID del reintegro al que se le asignan los documentos.
     *
     * Retorna:
     *   bool — True si la actualización afectó al menos una fila (éxito), False si no (reintegro no encontrado).
     *
     * Notas:
     *   - Actualiza también el campo `updated_at` con NOW().
     *   - Si la tabla o columnas tienen nombres distintos, ajustar la consulta.
    '''
    try:
        query = """
            UPDATE public.reintegro
            SET estado = 'ESPERANDO_ADJUNTOS',
                adjuntos_confirmados = FALSE,
                updated_at = NOW()
            WHERE reintegro_id = $1
        """
        result = await connection.execute(query, reintegro_id)

        # `execute` devuelve algo tipo 'UPDATE <n_rows>'
        if result and result.startswith("UPDATE"):
            try:
                updated = int(result.split()[1])
                return updated > 0
            except (IndexError, ValueError):
                # Resultado inesperado; considerar que no se actualizó
                return False
        return False
    except Exception as e:
        raise Exception(f"Error en utils.add_docs_reintegro: {e}")

'''
 * Lista reintegros para un afiliado dentro de un rango de fechas (inclusive).
 *
 * Parámetros:
 *   connection (asyncpg.Connection) — Conexión a la base de datos.
 *   afiliado_id (int) — ID del afiliado.
 *   fecha_desde (date) — Fecha inicial (inclusive).
 *   fecha_hasta (date) — Fecha final (inclusive).
 *
 * Retorna:
 *   List[dict] — Lista de reintegros (cada item es un dict con columnas seleccionadas).
 *
 * Notas:
 *   - Se usa fecha_presentacion::date para comparar solo la parte fecha.
'''
async def list_reintegros_por_afiliado_y_rango(
    connection: asyncpg.Connection,
    afiliado_id: int,
    fecha_desde: date,
    fecha_hasta: date,
) -> List[dict]:
    try:
        query = """
            SELECT reintegro_id, afiliado_id, estado, total_presentado, total_aprobado, fecha_presentacion, updated_at
            FROM public.reintegro
            WHERE afiliado_id = $1
              AND fecha_presentacion::date BETWEEN $2 AND $3
            ORDER BY fecha_presentacion DESC
        """
        rows = await connection.fetch(query, afiliado_id, fecha_desde, fecha_hasta)
        return [dict(r) for r in rows]
    except Exception as e:
        raise Exception(f"Error en utils.list_reintegros_por_afiliado_y_rango: {e}")

'''
Genera una nota para el reintegro y devuelve una URL de descarga accesible.
- Busca el afiliado por numero_afiliado en la tabla public.afiliado.
- Rellena la plantilla Word con los datos del afiliado y la descripcion.
- Guarda el archivo en una carpeta accesible y devuelve la URL de descarga.
'''
async def new_nota_reintegro(
    connection: asyncpg.Connection,
    descripcion: str,
    numero_afiliado: str,
) -> dict:
    from datetime import datetime
    from docx import Document
    import os

    try:
        # Buscar afiliado por numero_afiliado
        query_af = """
            SELECT nombre, apellido, domicilio, tel, email, tipo_doc, nro_doc, numero_afiliado
            FROM public.afiliado
            WHERE numero_afiliado = $1
            LIMIT 1
        """
        af = await connection.fetchrow(query_af, numero_afiliado)
        if not af:
            return {'exito': False, 'error': f'Afiliado no encontrado: {numero_afiliado}'}

        # Preparar datos del afiliado
        nombre_completo = f"{(af.get('nombre') or '').strip()} {(af.get('apellido') or '').strip()}".strip()
        domicilio = af.get('domicilio') or ''
        telefono = af.get('tel') or ''
        email = af.get('email') or ''
        carne_numero = af.get('numero_afiliado') or numero_afiliado
        documento_numero = af.get('nro_doc') or ''

        # Usar la fecha actual
        fecha_obj = datetime.now()
        meses = [
            'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
            'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
        ]
        fecha_dia = str(fecha_obj.day)
        fecha_mes = meses[fecha_obj.month - 1]
        fecha_anio = str(fecha_obj.year)

        # Usar solo la descripcion provista
        texto_solicitud = descripcion or ''

        # Cargar plantilla (ruta relativa simple)
        plantilla_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'static', 'plantilla_nota_reintegro.docx'))
        doc = Document(plantilla_path)

        # Reemplazos simples
        reemplazos = {
            '{{FECHA_DIA}}': fecha_dia,
            '{{FECHA_MES}}': fecha_mes,
            '{{FECHA_ANIO}}': fecha_anio,
            '{{TEXTO_SOLICITUD}}': texto_solicitud,
            '{{NOMBRE_COMPLETO}}': nombre_completo,
            '{{DOMICILIO}}': domicilio,
            '{{TELEFONO}}': telefono,
            '{{EMAIL}}': email,
            '{{CARNE_NUMERO}}': carne_numero,
            '{{DOCUMENTO_NUMERO}}': documento_numero
        }

        # Reemplazar en párrafos
        for parrafo in doc.paragraphs:
            for marcador, valor in reemplazos.items():
                if marcador in parrafo.text:
                    for run in parrafo.runs:
                        if marcador in run.text:
                            run.text = run.text.replace(marcador, str(valor))

        # Reemplazar en tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        for marcador, valor in reemplazos.items():
                            if marcador in parrafo.text:
                                for run in parrafo.runs:
                                    if marcador in run.text:
                                        run.text = run.text.replace(marcador, str(valor))

        # Guardar archivo en carpeta accesible públicamente
        nombre_archivo = f"Nota_OSEP_{carne_numero}_{fecha_obj.strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Carpeta pública para archivos descargables
        carpeta_publica = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'static', 'downloads'))
        os.makedirs(carpeta_publica, exist_ok=True)
        
        ruta_docx = os.path.join(carpeta_publica, nombre_archivo)
        doc.save(ruta_docx)

        # Generar URL de descarga accesible
        # NOTA: Ajustar el dominio según tu configuración del servidor
        download_url = f"/static/downloads/{nombre_archivo}"

        return {
            'exito': True,
            'archivo': nombre_archivo,
            'download_url': download_url,
            'mensaje': f'Nota generada correctamente. Puede descargarla desde: {download_url}'
        }

    except Exception as e:
        return {
            'exito': False,
            'error': str(e)
        }
